import os
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configurar o caminho do Tesseract
pytesseract.pytesseract_cmd = r'C:\Arquivos de Programas\Tesseract-OCR\tesseract.exe'

def pdf_to_images(pdf_path, output_folder):
    """Converte um PDF em imagens."""
    images = convert_from_path(pdf_path)
    image_paths = []
    for i, img in enumerate(images):
        img_path = os.path.join(output_folder, f'page_{i+1}.png')
        img.save(img_path, 'PNG')
        image_paths.append(img_path)
    return image_paths

def ocr_images_to_text(image_paths):
    """Extrai texto das imagens usando OCR e tenta preservar a formatação."""
    text_blocks = []
    for img_path in image_paths:
        text = pytesseract.image_to_string(img_path, lang='por', config='--psm 6')
        text_blocks.append(text.strip())
    return text_blocks

def save_text_to_word(text_blocks, output_docx):
    """Salva o texto extraído em um arquivo Word com formatação básica."""
    doc = Document()
    
    for text in text_blocks:
        if text.strip():
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)  # Tamanho do texto
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinhamento padrão
    
    doc.save(output_docx)

def clean_temp_folder(output_folder):
    """Remove os arquivos temporários da pasta temp_images."""
    for file in os.listdir(output_folder):
        file_path = os.path.join(output_folder, file)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Erro ao deletar {file_path}: {e}")

def main():
    # Solicitar entrada do usuário via variável de ambiente
    pdf_path = os.getenv("PDF_PATH", "arquivo_padrao.pdf")
    output_folder = "temp_images"
    output_docx = os.getenv("OUTPUT_DOCX", "resultado.docx")

    os.makedirs(output_folder, exist_ok=True)
    
    print("Convertendo PDF em imagens...")
    image_paths = pdf_to_images(pdf_path, output_folder)
    
    print("Extraindo texto das imagens...")
    text_blocks = ocr_images_to_text(image_paths)
    
    print("Salvando texto no Word com formatação...")
    save_text_to_word(text_blocks, output_docx)
    
    print("Limpando arquivos temporários...")
    clean_temp_folder(output_folder)
    
    print(f"Processo concluído! Arquivo salvo em {output_docx}")

# Executa o programa
if __name__ == "__main__":
    main()
